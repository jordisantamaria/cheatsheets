from docx import Document
document = Document()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
document.add_paragraph('3rth paragraph.')

document.save('test.docx')
